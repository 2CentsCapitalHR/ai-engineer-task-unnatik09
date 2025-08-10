from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _ensure_comments_part(doc):
    """
    Make sure the underlying package has a comments part we can append to.
    This tries to be robust across python-docx versions.
    """
    part = doc.part
    if not hasattr(part, "_comments"):
        # create a comments part container if absent
        # python-docx usually has _comments when comments.xml exists;
        # if not, create a minimal comments element. This is low-level and may
        # still fail on some setups. We catch exceptions in mark_file and fallback.
        comments = OxmlElement("w:comments")
        part._comments = type("C", (), {"_element": comments})
    return part

def _add_comment(paragraph, text, author="Corporate Agent", initials="AI"):
    """
    Try to insert a proper Word comment anchored to the paragraph.
    If any low-level manipulation fails, raise and let caller fallback to inline note.
    """
    doc = paragraph.part
    part = _ensure_comments_part(paragraph._p.getparent().part if False else paragraph.part)

    # fetch next id (create attribute if missing)
    if not hasattr(part, "_next_comment_id"):
        setattr(part, "_next_comment_id", 0)
    c_id = part._next_comment_id
    part._next_comment_id = c_id + 1

    # build <w:comment w:id="X" w:author="..." w:initials="..."><w:p>...</w:p></w:comment>
    comment = OxmlElement("w:comment")
    comment.set(qn("w:id"), str(c_id))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:initials"), initials)

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    comment.append(p)

    # Append comment element to comments container
    part._comments._element.append(comment)

    # Insert commentRangeStart before paragraph text, commentRangeEnd after paragraph text,
    # and a commentReference in a new run so Word shows the comment marker.
    # We'll attach at the start of the paragraph run.
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(c_id))
    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(c_id))
    ref_r = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(c_id))
    ref_r.append(ref)

    # Add to paragraph xml
    p_elm = paragraph._p
    # insert start before first child, append end and reference at the end
    p_elm.insert(0, start)
    p_elm.append(end)
    p_elm.append(ref_r)

def _append_inline_note(paragraph, text):
    """
    Fallback if comment insertion fails — append a clearly-delimited reviewer note.
    """
    note = paragraph._parent.add_paragraph()
    note.add_run(f"[REVIEWER NOTE] {text}")

def mark_file(src, dst, findings):
    """
    Mark a .docx file with comments for each finding.
    `findings` is a list of dicts with keys:
      - pattern (string to search; already normalized lower)
      - issue (short description)
      - cite (legal citation string)
      - suggest (suggested replacement text)
    The function tries to add true Word comments; on failure it appends inline notes.
    """
    doc = Document(src)

    for f in findings:
        matched = False
        pattern = f.get("pattern", "").lower()
        note_text = f'{f.get("issue")} – {f.get("cite")}\nSuggestion: {f.get("suggest")}'
        for p in doc.paragraphs:
            if pattern and pattern in p.text.lower():
                try:
                    _add_comment(p, note_text)
                except Exception:
                    # fallback
                    _append_inline_note(p, note_text)
                matched = True
                break
        if not matched:
            # If no paragraph matched, append note to document end
            try:
                last = doc.paragraphs[-1]
                _add_comment(last, note_text)
            except Exception:
                _append_inline_note(doc.paragraphs[-1], note_text)

    doc.save(dst)
